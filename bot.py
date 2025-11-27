import os
import asyncio
import logging
from datetime import datetime, timedelta
from collections import defaultdict
import re

import openai
from dotenv import load_dotenv
from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart
from aiogram.types import (
    Message,
    CallbackQuery,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    LabeledPrice,
    PreCheckoutQuery,
    FSInputFile,
)

from docx import Document
import PyPDF2

load_dotenv()

BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PAYMENT_PROVIDER_TOKEN = os.getenv("PAYMENT_PROVIDER_TOKEN")

openai.api_key = OPENAI_API_KEY

logging.basicConfig(level=logging.INFO)

bot = Bot(BOT_TOKEN)
dp = Dispatcher()

FREE_LIMIT = 10

PRICE_INDIV = 199
PRICE_PACK5 = 75
PRICE_PACK10 = 129
PRICE_PACK20 = 239

PRICE_DOC_COMPOSE = 75
PRICE_DOC_CHECK = 99
PRICE_PLAN = 149
PRICE_FULL_SUPPORT = 299

users = defaultdict(
    lambda: {
        "free_left": FREE_LIMIT,
        "paid_left": 0,
        "last_reset": datetime.now(),
        "consult_active": False,
        "service": None,
        "doc_format": None,
        "case_mode": None,
        "case_summary": None,
    }
)

SYSTEM_PROMPT_BASE = """
Ты — «Адвокат X», онлайн-юрист по праву РФ и опытный процессуалист.

ВАЖНО:
• Ты работаешь как ИИ-помощник. Ты не являешься адвокатом, не состоишь в палате и не оказываешь официальную юридическую помощь.
• Все ответы носят справочный и информационный характер. Пользователь сам несёт ответственность за свои решения.
• Ты строго работаешь в рамках законодательства РФ. Не предлагаешь серые, полулегальные, криминальные схемы, обход закона, фиктивные документы и т.п.
• Если пользователь прямо просит помочь нарушить закон, обойти блокировки, налоговые/призывные/иные обязательства незаконным способом — мягко отказываешь и предлагаешь только легальные пути.

СТИЛЬ РАБОТЫ:
1) Всегда уточняешь недостающие факты, если без них нельзя дать точный ответ (но при этом даёшь общую рамку, если это возможно).
2) Строишь ответы по чёткой структуре (если формат вопроса позволяет):
   - Краткий вывод по ситуации простым языком;
   - Правовое основание (важные нормы, уровни: Конституция, кодексы, законы, подзаконка — если уместно);
   - Варианты действий (минимум 2–3 сценария, если есть выбор);
   - Риски и подводные камни;
   - Пошаговый план действий.
3) Не предлагаешь ничего, что может очевидно ухудшить положение пользователя без предупреждения. О рисках говоришь прямо.
4) Даёшь максимально выгодную стратегию в рамках закона, исходя из интересов пользователя (время, деньги, нервы).
5) Если фактов мало — вначале коротко поясняешь общую ситуацию, а затем перечисляешь, что нужно уточнить у пользователя.
6) Пишешь человеческим языком, без канцелярита, но с сохранением юридической точности. Расшифровываешь сложные термины простыми словами.

Ответы делай структурированными, читаемыми, с подзаголовками и нумерацией там, где это уместно.
"""

def reset_limits(user_id: int):
    u = users[user_id]
    now = datetime.now()
    if now - u["last_reset"] >= timedelta(days=1):
        u["free_left"] = FREE_LIMIT
        u["last_reset"] = now


async def ask_model(messages: list) -> str:
    def _call():
        return openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.2,
        )

    r = await asyncio.to_thread(_call)
    return r["choices"][0]["message"]["content"].strip()


async def ask_short_consult(text: str) -> str:
    return await ask_model(
        [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {
                "role": "user",
                "content": (
                    "Режим: краткая консультация.\n"
                    "Задача: быстро сориентировать человека по его ситуации.\n\n"
                    "Сделай короткий, но ёмкий ответ (примерно 5–8 предложений):\n"
                    "• дай общий вывод по ситуации простым языком;\n"
                    "• обозначь, какие законы или области права здесь задействованы (без перегруза статьями);\n"
                    "• укажи, что человек может сделать прямо сейчас (2–3 первых шага);\n"
                    "• если явно не хватает фактов — в конце коротко напиши, что стоит уточнить.\n\n"
                    "Не используй сложные конструкции, пиши как живой онлайн-юрист в чате.\n\n"
                    f"ВОПРОС / СИТУАЦИЯ:\n{text}"
                ),
            },
        ]
    )


async def ask_individual_consult(text: str) -> str:
    return await ask_model(
        [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {
                "role": "user",
                "content": (
                    "Режим: углубленный разбор одного вопроса.\n"
                    "Представь, что клиент пришёл на очную консультацию к практикующему юристу.\n\n"
                    "Сделай развёрнутый профессиональный разбор ситуации с чёткой структурой:\n"
                    "1) Краткий итог (что происходит и как это выглядит с точки зрения права).\n"
                    "2) Нормы права РФ, которые применимы (укажи основные статьи/акты, но без перегруза).\n"
                    "3) Варианты стратегии и тактики:\n"
                    "   • мягкий/мирный вариант;\n"
                    "   • более жёсткий/конфликтный вариант;\n"
                    "   • альтернативы, если такие есть.\n"
                    "4) Практические советы, как усилить позицию клиента (доказательства, документы, свидетели, сроки и т.п.).\n"
                    "5) Риски: чем каждый вариант может обернуться, на что обратить внимание.\n"
                    "6) Итоговый рекомендуемый курс действий от тебя.\n\n"
                    "Стиль — как у внимательного онлайн-юриста: понятно, структурировано, без серых/нелегальных схем.\n\n"
                    f"СИТУАЦИЯ КЛИЕНТА:\n{text}"
                ),
            },
        ]
    )


async def ask_doc_compose(text: str) -> str:
    return await ask_model(
        [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {
                "role": "user",
                "content": (
                    "Режим: составление юридического документа.\n"
                    "Нужно подготовить максимально готовый к использованию текст документа по описанной ситуации.\n\n"
                    "Требования к документу:\n"
                    "• сам выбери оптимальный тип документа (претензия, заявление, жалоба, иск, объяснительная, договор, соглашение и т.п.), исходя из интересов клиента и закона;\n"
                    "• используй деловой, официальный стиль, без эмоциональных оценок;\n"
                    "• не включай вымышленные реквизиты (ФИО, адрес, ИНН, номера договоров и т.п.) — оставь шаблонные места.\n\n"
                    "Структура документа (универсальный шаблон, адаптируй под ситуацию):\n"
                    "1) «Шапка» (кому, от кого — шаблонно, без конкретных данных клиента);\n"
                    "2) Вводная часть: кратко и по фактам опиши ситуацию (кто, с кем, когда, какие отношения, какие документы есть);\n"
                    "3) Правовое обоснование: ссылки на нормы права РФ, если это уместно и полезно;\n"
                    "4) Требования / просьба / условия (чего добивается заявитель);\n"
                    "5) Заключительная часть: фразы про сроки, порядок ответа, приложения (если логично);\n"
                    "6) Места для даты и подписи.\n\n"
                    "Не добавляй в текст указания на то, что это сгенерировал ИИ. Документ должен выглядеть как нормальный человеческий документ.\n\n"
                    f"СИТУАЦИЯ ДЛЯ ДОКУМЕНТА:\n{text}"
                ),
            },
        ]
    )


async def ask_doc_check(text: str) -> str:
    return await ask_model(
        [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {
                "role": "user",
                "content": (
                    "Режим: проверка юридического документа.\n"
                    "Ниже дан текст документа или его черновика.\n\n"
                    "Твоя задача — провести юридический разбор и дать рекомендации по усилению документа. Структурируй ответ так:\n"
                    "1) Общая оценка документа (понятно ли, логично ли, каких целей он достигает).\n"
                    "2) Потенциальные риски и слабые места:\n"
                    "   • двусмысленные формулировки;\n"
                    "   • пункты, которые могут быть использованы против клиента;\n"
                    "   • нарушения или возможные проблемы с точки зрения права.\n"
                    "3) Рекомендации по усилению документа:\n"
                    "   • какие формулировки лучше заменить и на какие;\n"
                    "   • какие пункты стоит добавить (сроки, ответственность, порядок расторжения, подсудность и т.п.);\n"
                    "   • на что обратить внимание перед подписанием.\n"
                    "4) Если документ в принципе нормальный — коротко напиши об этом и перечисли только важные доработки.\n\n"
                    "Не переписывай весь документ заново, сосредоточься на анализе и ключевых правках.\n\n"
                    f"ТЕКСТ ДОКУМЕНТА ДЛЯ ПРОВЕРКИ:\n{text}"
                ),
            },
        ]
    )


async def ask_plan_actions(text: str) -> str:
    return await ask_model(
        [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {
                "role": "user",
                "content": (
                    "Режим: составление пошагового плана действий по делу.\n"
                    "Нужно помочь человеку понять, КАК КОНКРЕТНО ему двигаться дальше.\n\n"
                    "Структура ответа:\n"
                    "1) Краткий вывод по ситуации (что в целом происходит с точки зрения права).\n"
                    "2) Общая стратегия (какой результат мы хотим получить в идеале).\n"
                    "3) Подробный пошаговый план (пункты 1, 2, 3, ...), с примерными сроками:\n"
                    "   • куда обращаться (суд, госорган, полиция, работодатель, контрагент и т.п.);\n"
                    "   • что именно подать или подписать (заявление, жалоба, претензия, иск и т.д.);\n"
                    "   • что сделать ДО обращения (собрать доказательства, переписку, чеки и т.п.).\n"
                    "4) Какие документы и доказательства собирать (по пунктам).\n"
                    "5) Возможные риски и альтернативные варианты, если что-то пойдёт не по плану.\n\n"
                    "Пиши так, чтобы человек мог буквально идти по пунктам, как по чек-листу.\n\n"
                    f"СИТУАЦИЯ КЛИЕНТА:\n{text}"
                ),
            },
        ]
    )


async def ask_full_support(text: str, case_summary: str | None = None) -> str:
    case_part = (
        f"Краткое описание дела, которое уже ведётся: {case_summary}\n\n"
        if case_summary
        else ""
    )
    return await ask_model(
        [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {
                "role": "user",
                "content": (
                    "Режим: полное сопровождение одного дела.\n"
                    "Представь, что ты сопровождаешь одно конкретное дело клиента на протяжении времени.\n"
                    "Каждое новое сообщение — это обновления по делу, вопросы, новые документы, ответы оппонентов и т.п.\n\n"
                    f"{case_part}"
                    "Сделай детальный разбор текущего сообщения и предложи тактику на ближайший шаг:\n"
                    "1) Что означает то, что происходит сейчас (с точки зрения права и практики).\n"
                    "2) Как это влияет на общую стратегию по делу.\n"
                    "3) Что клиенту делать дальше (1–3 первых шага, максимально конкретно).\n"
                    "4) Какие риски у этих шагов и как их можно смягчить.\n"
                    "5) Если нужно — какие документы/доказательства подготовить или скорректировать.\n\n"
                    "Всегда исходи из того, что дело одно и ты знаешь контекст из предыдущих сообщений.\n\n"
                    f"ТЕКУЩЕЕ СООБЩЕНИЕ КЛИЕНТА:\n{text}"
                ),
            },
        ]
    )


def kb_main():
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="Задать вопрос", callback_data="mode_basic")],
            [InlineKeyboardButton(text="Каталог услуг", callback_data="catalog")],
            [InlineKeyboardButton(text="Пакеты сообщений", callback_data="menu")],
            [InlineKeyboardButton(text="ℹ️ Инфо", callback_data="info")],
            [InlineKeyboardButton(text="⚖️ Условия использования", callback_data="terms")],
        ]
    )


def kb_catalog():
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Углубленный разбор — {PRICE_INDIV} ₽", callback_data="srv_indiv"
                )
            ],
            [
                InlineKeyboardButton(
                    text=f"Составить документ — {PRICE_DOC_COMPOSE} ₽",
                    callback_data="srv_doc_compose",
                )
            ],
            [
                InlineKeyboardButton(
                    text=f"Проверить документ — {PRICE_DOC_CHECK} ₽",
                    callback_data="srv_doc_check",
                )
            ],
            [
                InlineKeyboardButton(
                    text=f"План действий — {PRICE_PLAN} ₽", callback_data="srv_plan"
                )
            ],
            [
                InlineKeyboardButton(
                    text=f"Полное сопровождение дела — {PRICE_FULL_SUPPORT} ₽",
                    callback_data="srv_full_support",
                )
            ],
            [InlineKeyboardButton(text="⬅️ Назад", callback_data="back_main")],
        ]
    )


def kb_menu():
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"5 сообщений — {PRICE_PACK5} ₽", callback_data="buy5"
                )
            ],
            [
                InlineKeyboardButton(
                    text=f"10 сообщений — {PRICE_PACK10} ₽", callback_data="buy10"
                )
            ],
            [
                InlineKeyboardButton(
                    text=f"20 сообщений — {PRICE_PACK20} ₽", callback_data="buy20"
                )
            ],
            [
                InlineKeyboardButton(
                    text="⬅️ Назад в главное меню", callback_data="back_main"
                )
            ],
        ]
    )


def kb_main_button(doc: bool = False):
    if not doc:
        return InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")]
            ]
        )
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"📄 Составить документ по этой ситуации — {PRICE_DOC_COMPOSE} ₽",
                    callback_data="buy_doc_compose",
                )
            ],
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")],
        ]
    )


def kb_full_support():
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Составить документ по этой ситуации — {PRICE_DOC_COMPOSE} ₽",
                    callback_data="buy_doc_compose",
                )
            ],
            [InlineKeyboardButton(text="✅ Дело решено", callback_data="case_done")],
            [InlineKeyboardButton(text="Пакеты сообщений", callback_data="menu")],
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")],
        ]
    )


async def create_invoice(chat_id, title, description, payload, amount_rub: int):
    amount_cents = int(amount_rub) * 100
    prices = [LabeledPrice(label=title, amount=amount_cents)]

    await bot.send_invoice(
        chat_id=chat_id,
        title=title,
        description=description,
        payload=payload,
        provider_token=PAYMENT_PROVIDER_TOKEN,
        currency="RUB",
        prices=prices,
        need_email=True,
        send_email_to_provider=True,
    )


def make_docx_file(text: str, uid: int) -> str:
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    filename = f"document_{uid}_{int(datetime.now().timestamp())}.docx"
    doc.save(filename)
    return filename


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if ext == ".pdf":
        text = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)
    return ""


def is_same_case(new_text: str, base_text: str | None) -> bool:
    if not base_text:
        return True

    def tokens(s: str):
        return {w for w in re.findall(r"\w+", s.lower()) if len(w) >= 4}

    a = tokens(new_text)
    b = tokens(base_text)
    if not a or not b:
        return True
    inter = len(a & b)
    ratio = inter / min(len(a), len(b))
    return ratio >= 0.25


def need_doc_button(answer: str) -> bool:
    a = answer.lower()
    return bool(
        re.search(
            r"заявлен|претензи|жалоб|иск|договор|расписк|соглашени|ходатайств|акт|протокол",
            a,
        )
    )


@dp.message(CommandStart())
async def start_cmd(message: Message):
    uid = message.from_user.id
    reset_limits(uid)
    u = users[uid]
    await message.answer(
        "Здравствуйте! Я — «Адвокат X», юридический ИИ-помощник по праву РФ.\n\n"
        "Что я могу для вас сделать:\n"
        "• ответить на ваш юридический вопрос простым языком;\n"
        "• подсказать варианты стратегии в споре;\n"
        "• помочь с текстом юридического документа;\n"
        "• помочь выстроить план действий по делу.\n\n"
        "Важно:\n"
        "• Я работаю на базе нейросети и не являюсь адвокатом или официальным представителем.\n"
        "• Все ответы носят информационный характер и не заменяют очную консультацию.\n"
        "• Я не помогаю нарушать закон или действовать «в серую».\n\n"
        f"Бесплатных сообщений на сегодня: {u['free_left']}\n"
        f"Оплаченных сообщений: {u['paid_left']}\n\n"
        "Опишите вашу ситуацию одним сообщением или выберите нужный раздел ниже:",
        reply_markup=kb_main(),
    )


@dp.callback_query(F.data == "back_main")
async def back_main(call: CallbackQuery):
    await call.message.answer("Главное меню:", reply_markup=kb_main())
    await call.answer()


@dp.callback_query(F.data == "catalog")
async def catalog(call: CallbackQuery):
    await call.message.answer(
        "Каталог услуг:\n\n"
        "Здесь собраны расширенные форматы работы, если вам недостаточно короткого ответа в чате.\n"
        "Выберите услугу, почитайте описание и, если подходит, оформите её.",
        reply_markup=kb_catalog(),
    )
    await call.answer()


@dp.callback_query(F.data == "mode_basic")
async def mode_basic(call: CallbackQuery):
    await call.message.answer(
        "Режим «Задать вопрос» активирован.\n\n"
        "Напишите одним сообщением вашу ситуацию или вопрос. Я дам краткий, но по существу ответ:\n"
        "• что происходит с точки зрения права;\n"
        "• с чего лучше начать;\n"
        "• что учесть, чтобы не сделать хуже.\n\n"
        "Если понадобится более глубокий разбор, вы всегда сможете выбрать услугу из каталога."
    )
    await call.answer()


@dp.callback_query(F.data == "menu")
async def menu(call: CallbackQuery):
    await call.message.answer(
        "Пакеты сообщений:\n\n"
        "Если вам не хватает бесплатного дневного лимита, вы можете докупить сообщения.\n"
        "Они расходуются на все краткие консультации, когда бесплатный лимит уже исчерпан.",
        reply_markup=kb_menu(),
    )
    await call.answer()


@dp.callback_query(F.data == "info")
async def info(call: CallbackQuery):
    text = (
        "ℹ️ Кратко об услугах бота «Адвокат X»:\n\n"
        "1. Задать вопрос (бесплатный лимит)\n"
        "   Краткий ответ по вашей ситуации. Подходит, если вы хотите:\n"
        "   • понять, есть ли вообще смысл что-то делать;\n"
        "   • быстро сориентироваться, к какому органу идти;\n"
        "   • получить первые шаги без глубокого разбора.\n\n"
        f"2. Углубленный разбор — {PRICE_INDIV} ₽\n"
        "   Подробная консультация по одному вопросу:\n"
        "   • разберём факты и документы;\n"
        "   • рассмотрим несколько стратегий поведения;\n"
        "   • обсудим риски, сроки и доказательства;\n"
        "   • в конце получите ясный план, как действовать.\n\n"
        f"3. Составить документ — {PRICE_DOC_COMPOSE} ₽\n"
        "   Подготовка текста юридического документа под вашу ситуацию:\n"
        "   • претензия, заявление, жалоба, иск, договор и др.;\n"
        "   • деловой стиль и понятная структура;\n"
        "   • вам останется только подставить свои данные и распечатать.\n\n"
        f"4. Проверить документ — {PRICE_DOC_CHECK} ₽\n"
        "   Анализ уже готового документа:\n"
        "   • нахожу слабые места и потенциальные риски;\n"
        "   • подсказываю, что может быть использовано против вас;\n"
        "   • предлагаю, как переформулировать пункты, чтобы усилить вашу позицию.\n\n"
        f"5. План действий — {PRICE_PLAN} ₽\n"
        "   Подробный чек-лист по шагам:\n"
        "   • что делать сначала, что потом;\n"
        "   • куда идти, какие документы готовить;\n"
        "   • как реагировать на типичные ответы и отказы.\n\n"
        f"6. Полное сопровождение дела — {PRICE_FULL_SUPPORT} ₽\n"
        "   Работа по одному делу с учётом всех обновлений:\n"
        "   • вы описываете ситуацию, дальше дополняете её по мере развития событий;\n"
        "   • я помогаю скорректировать стратегию, отвечать на письма, готовиться к шагам;\n"
        "   • удобно, если дело длится долго и всё меняется.\n\n"
        "Оплата проходит через Telegram (ЮKassa). Чек приходит на ваш e-mail.\n"
        "Помните: ответы носят информационный характер и не являются официальной юридической консультацией."
    )
    await call.message.answer(text)
    await call.answer()


@dp.callback_query(F.data == "terms")
async def terms(call: CallbackQuery):
    text = (
        "⚖️ Условия использования:\n\n"
        "• Бот «Адвокат X» работает на базе нейросети и не является адвокатом, юристом или представительством.\n"
        "• Ответы носят справочный и информационный характер и не являются официальной юридической "
        "консультацией, заключением или публичной офертой.\n"
        "• Ответ формируется автоматически на основе введённых вами данных. Вы самостоятельно принимаете решения "
        "и несёте ответственность за их последствия.\n"
        "• Запрещено использовать бота для подготовки явно незаконных действий, обхода закона, подделки документов, "
        "обмана органов, уклонения от ответственности и т.п.\n"
        "• Сервис предоставляется «как есть», без гарантий результата в суде или перед госорганами.\n"
    )
    await call.message.answer(text)
    await call.answer()


@dp.callback_query(F.data == "srv_indiv")
async def srv_indiv(call: CallbackQuery):
    text = (
        "Услуга: «Углубленный разбор».\n\n"
        "Для кого подходит:\n"
        "• у вас есть конкретная проблема (с работодателем, арендодателем, банком, соседями, госорганом и т.п.);\n"
        "• нужно детально понять, на чьей стороне закон и какие варианты действий у вас есть;\n"
        "• вы хотите минимизировать риски и заранее понимать возможные последствия.\n\n"
        "Что вы получаете:\n"
        "• подробный разбор ваших фактов и документов;\n"
        "• объяснение, какие нормы закона применимы в вашей ситуации;\n"
        "• несколько сценариев действий (от мягкого до более жёсткого подхода);\n"
        "• рекомендации по доказательствам и ключевым шагам;\n"
        "• итоговый курс действий от «Адвоката X».\n\n"
        "Услуга рассчитана на один вопрос/одну ситуацию и один развёрнутый ответ."
    )
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Оплатить — {PRICE_INDIV} ₽", callback_data="buy_indiv"
                )
            ],
            [InlineKeyboardButton(text="⬅️ Назад к каталогу", callback_data="catalog")],
        ]
    )
    await call.message.answer(text, reply_markup=kb)
    await call.answer()


@dp.callback_query(F.data == "srv_doc_compose")
async def srv_doc_compose(call: CallbackQuery):
    text = (
        "Услуга: «Составить документ».\n\n"
        "Для чего это нужно:\n"
        "• вы хотите направить претензию, жалобу, заявление, иск или иной документ, но не знаете, как его грамотно оформить;\n"
        "• боитесь написать «на эмоциях» и хотите деловой, спокойный и юридически выверенный текст;\n"
        "• вам нужно, чтобы документ выглядел так, будто его готовил живой юрист.\n\n"
        "Что вы получаете:\n"
        "• готовый текст юридического документа под вашу ситуацию;\n"
        "• понятную структуру (шапка, факты, правовое обоснование, требования, заключение);\n"
        "• деловой стиль, без лишних эмоций и воды;\n"
        "• формат DOCX, который удобно редактировать и отправлять.\n\n"
        "После оплаты вам нужно будет подробно описать ситуацию и указать, какой именно документ вы хотите получить."
    )
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Оплатить — {PRICE_DOC_COMPOSE} ₽",
                    callback_data="buy_doc_compose",
                )
            ],
            [InlineKeyboardButton(text="⬅️ Назад к каталогу", callback_data="catalog")],
        ]
    )
    await call.message.answer(text, reply_markup=kb)
    await call.answer()


@dp.callback_query(F.data == "srv_doc_check")
async def srv_doc_check(call: CallbackQuery):
    text = (
        "Услуга: «Проверить документ».\n\n"
        "Для чего это нужно:\n"
        "• вы уже подготовили договор, претензию, заявление, расписку или другой документ;\n"
        "• хотите понять, нет ли там подводных камней и слабых формулировок;\n"
        "• переживаете, что текст можно будет повернуть против вас.\n\n"
        "Что вы получаете:\n"
        "• анализ текста документа с точки зрения интересов клиента;\n"
        "• указание рисков и спорных формулировок;\n"
        "• предложения, как изменить текст, чтобы усилить вашу позицию;\n"
        "• рекомендации, какие пункты стоит добавить (сроки, ответственность, порядок расторжения и т.п.).\n\n"
        "Документ можно отправить текстом или файлом (PDF, DOCX, TXT). Фото лучше не использовать — текст с них я не читаю."
    )
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Оплатить — {PRICE_DOC_CHECK} ₽",
                    callback_data="buy_doc_check",
                )
            ],
            [InlineKeyboardButton(text="⬅️ Назад к каталогу", callback_data="catalog")],
        ]
    )
    await call.message.answer(text, reply_markup=kb)
    await call.answer()


@dp.callback_query(F.data == "srv_plan")
async def srv_plan(call: CallbackQuery):
    text = (
        "Услуга: «План действий».\n\n"
        "Для кого подходит:\n"
        "• вы запутались в шагях: куда идти, что писать, кому звонить;\n"
        "• не хотите тратить время и нервы на хаотичные попытки.\n\n"
        "Что вы получаете:\n"
        "• чёткий по шагам план (как чек-лист);\n"
        "• указание, куда обращаться и с чем именно (заявления, жалобы, документы);\n"
        "• примерные сроки и последовательность действий;\n"
        "• предупреждения о типичных ошибках и рисках.\n\n"
        "После оплаты опишите вашу ситуацию максимально подробно — на основе этого я соберу для вас план."
    )
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Оплатить — {PRICE_PLAN} ₽", callback_data="buy_plan"
                )
            ],
            [InlineKeyboardButton(text="⬅️ Назад к каталогу", callback_data="catalog")],
        ]
    )
    await call.message.answer(text, reply_markup=kb)
    await call.answer()


@dp.callback_query(F.data == "srv_full_support")
async def srv_full_support(call: CallbackQuery):
    text = (
        "Услуга: «Полное сопровождение дела».\n\n"
        "Для чего это нужно:\n"
        "• ваше дело длится не один день (переписка, ответы, отказы, новые документы);\n"
        "• вы хотите, чтобы один и тот же онлайн-юрист помогал вам выстраивать стратегию на каждом шаге;\n"
        "• вам важно не просто разово спросить, а двигаться по делу последовательно.\n\n"
        "Что вы получаете:\n"
        "• сопровождение ОДНОГО конкретного дела;\n"
        "• помощь в выборе стратегии и корректировке тактики по мере развития ситуации;\n"
        "• подсказки, как реагировать на письма, ответы, претензии, действия оппонента или органов;\n"
        "• рекомендации по документам, доказательствам и следующими шагам.\n\n"
        "Если у вас начнётся новое отдельное дело, для него нужно будет оформить новую услугу."
    )
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=f"Оплатить — {PRICE_FULL_SUPPORT} ₽",
                    callback_data="buy_full_support",
                )
            ],
            [InlineKeyboardButton(text="⬅️ Назад к каталогу", callback_data="catalog")],
        ]
    )
    await call.message.answer(text, reply_markup=kb)
    await call.answer()


@dp.callback_query(F.data == "buy_indiv")
async def buy_indiv(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Углубленный разбор",
        "Развёрнутый разбор одной юридической ситуации с вариантами стратегии и рисков.",
        "individual",
        PRICE_INDIV,
    )
    await call.answer()


@dp.callback_query(F.data == "buy_doc_compose")
async def buy_doc_compose(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Составление документа",
        "Подготовка текста юридического документа под вашу ситуацию (формат DOCX).",
        "doc_compose",
        PRICE_DOC_COMPOSE,
    )
    await call.answer()


@dp.callback_query(F.data == "buy_doc_check")
async def buy_doc_check(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Проверка документа",
        "Проверка текста документа: риски, слабые места, рекомендации по усилению.",
        "doc_check",
        PRICE_DOC_CHECK,
    )
    await call.answer()


@dp.callback_query(F.data == "buy_plan")
async def buy_plan(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "План действий",
        "Подробный по шагам план действий по вашему юридическому вопросу.",
        "plan",
        PRICE_PLAN,
    )
    await call.answer()


@dp.callback_query(F.data == "buy_full_support")
async def buy_full_support(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Полное сопровождение дела",
        "Сопровождение одного дела: стратегия, тактика, помощь на каждом шаге.",
        "full_support",
        PRICE_FULL_SUPPORT,
    )
    await call.answer()


@dp.callback_query(F.data == "buy5")
async def buy5(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Пакет 5 сообщений",
        "Дополнительно 5 сообщений сверх бесплатного лимита.",
        "pack5",
        PRICE_PACK5,
    )
    await call.answer()


@dp.callback_query(F.data == "buy10")
async def buy10(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Пакет 10 сообщений",
        "Дополнительно 10 сообщений сверх бесплатного лимита.",
        "pack10",
        PRICE_PACK10,
    )
    await call.answer()


@dp.callback_query(F.data == "buy20")
async def buy20(call: CallbackQuery):
    uid = call.from_user.id
    await create_invoice(
        uid,
        "Пакет 20 сообщений",
        "Дополнительно 20 сообщений сверх бесплатного лимита.",
        "pack20",
        PRICE_PACK20,
    )
    await call.answer()


@dp.callback_query(F.data == "case_done")
async def case_done(call: CallbackQuery):
    uid = call.from_user.id
    u = users[uid]
    u["consult_active"] = False
    u["case_mode"] = None
    u["case_summary"] = None
    await call.message.answer(
        "Отмечаю дело как завершённое.\n"
        "Если появится новое отдельное дело, вы можете оформить новую услугу сопровождения.",
        reply_markup=kb_main(),
    )
    await call.answer()


@dp.pre_checkout_query()
async def checkout(q: PreCheckoutQuery):
    await bot.answer_pre_checkout_query(q.id, ok=True)


@dp.message(F.successful_payment)
async def paid(message: Message):
    uid = message.from_user.id
    payload = message.successful_payment.invoice_payload
    u = users[uid]

    if payload == "individual":
        u["service"] = "deep"
        await message.answer(
            "Услуга «Углубленный разбор» активирована.\n\n"
            "Пожалуйста, опишите вашу ситуацию максимально подробно: кто с кем, какие отношения, что уже произошло, "
            "какие документы есть и чего вы хотите добиться.",
            reply_markup=kb_main_button(),
        )
    elif payload == "doc_compose":
        u["service"] = "doc_compose"
        u["doc_format"] = "docx"
        await message.answer(
            "Оплата прошла.\n\n"
            "Опишите подробно вашу ситуацию и укажите, какой документ вы хотите получить "
            "(например: претензия, заявление, жалоба, иск, договор и т.п.). "
            "Чем подробнее вы опишете факты, тем точнее получится документ.",
            reply_markup=kb_main_button(),
        )
    elif payload == "doc_check":
        u["service"] = "doc_check"
        await message.answer(
            "Оплата прошла.\n\n"
            "Пришлите документ для проверки:\n"
            "• текстом (скопируйте его сюда), или\n"
            "• файлом PDF / DOCX / TXT.\n\n"
            "Если пришлёте фото, я попрошу вас отправить текст вручную, так как не читаю текст с изображений.",
            reply_markup=kb_main_button(),
        )
    elif payload == "plan":
        u["service"] = "plan"
        await message.answer(
            "Оплата прошла.\n\n"
            "Опишите вашу ситуацию так, как будто рассказываете её юристу: кто участники, что уже произошло, "
            "какие есть документы, ответы, сроки и чего вы хотите добиться. На основе этого я сделаю для вас план действий.",
            reply_markup=kb_main_button(),
        )
    elif payload == "full_support":
        u["consult_active"] = True
        u["case_mode"] = "full"
        u["case_summary"] = None
        await message.answer(
            "Полное сопровождение дела активировано.\n\n"
            "Сначала подробно опишите вашу ситуацию (это будет базовое описание дела). "
            "Дальше вы сможете дополнять её новыми сообщениями по мере развития событий, а я буду помогать "
            "корректировать стратегию и следующие шаги.",
            reply_markup=kb_main_button(),
        )
    elif payload == "pack5":
        u["paid_left"] += 5
        await message.answer(
            "Оплата прошла. К вашему лимиту добавлено 5 сообщений.",
            reply_markup=kb_main_button(),
        )
    elif payload == "pack10":
        u["paid_left"] += 10
        await message.answer(
            "Оплата прошла. К вашему лимиту добавлено 10 сообщений.",
            reply_markup=kb_main_button(),
        )
    elif payload == "pack20":
        u["paid_left"] += 20
        await message.answer(
            "Оплата прошла. К вашему лимиту добавлено 20 сообщений.",
            reply_markup=kb_main_button(),
        )


@dp.message(F.document)
async def doc_message(message: Message):
    uid = message.from_user.id
    u = users[uid]

    if u["service"] != "doc_check":
        await message.answer(
            "Сейчас проверка документа у вас не активирована.\n\n"
            "Если хотите проверить документ, откройте «Каталог услуг» и выберите «Проверить документ».",
            reply_markup=kb_main_button(),
        )
        return

    os.makedirs("files", exist_ok=True)
    file = message.document
    local_path = os.path.join("files", f"{uid}_{file.file_name}")
    await bot.download(file, destination=local_path)

    text = extract_text_from_file(local_path)
    if not text.strip():
        await message.answer(
            "Не удалось прочитать текст из файла.\n\n"
            "Поддерживаются форматы: PDF, DOCX, TXT.\n"
            "Попробуйте отправить документ в одном из этих форматов или скопируйте текст прямо в сообщение.",
            reply_markup=kb_main_button(),
        )
        return

    ans = await ask_doc_check(text)
    u["service"] = None
    await message.answer(ans, reply_markup=kb_main_button())


@dp.message(F.photo)
async def photo_message(message: Message):
    uid = message.from_user.id
    u = users[uid]

    if u["service"] == "doc_check":
        await message.answer(
            "Я пока не умею надёжно читать текст с фотографий.\n\n"
            "Пожалуйста, отправьте документ в виде текста или файла PDF / DOCX / TXT.",
            reply_markup=kb_main_button(),
        )
    else:
        await message.answer(
            "Если вы хотите проверить документ, сначала выберите услугу «Проверить документ» в каталоге и оплатите её.\n"
            "После этого отправьте текст или файл документа.",
            reply_markup=kb_main_button(),
        )


@dp.message(F.text)
async def msg(message: Message):
    uid = message.from_user.id
    text = message.text
    u = users[uid]

    reset_limits(uid)

    if u["service"] == "doc_compose":
        ans = await ask_doc_compose(text)
        await message.answer(ans)
        docx_path = make_docx_file(ans, uid)
        await message.answer_document(
            FSInputFile(docx_path),
            caption="Ваш документ в формате DOCX.",
        )
        u["service"] = None
        u["doc_format"] = None
        await message.answer(
            "Документ готов.\n"
            "Проверьте текст, при необходимости подкорректируйте детали и подставьте свои данные.",
            reply_markup=kb_main_button(),
        )
        return

    if u["service"] == "doc_check":
        ans = await ask_doc_check(text)
        u["service"] = None
        await message.answer(ans, reply_markup=kb_main_button())
        return

    if u["service"] == "plan":
        ans = await ask_plan_actions(text)
        u["service"] = None
        await message.answer(ans, reply_markup=kb_main_button())
        return

    if u["service"] == "deep":
        ans = await ask_individual_consult(text)
        doc_btn = need_doc_button(ans)
        u["service"] = None
        await message.answer(ans, reply_markup=kb_main_button(doc_btn))
        return

    if u["consult_active"] and u.get("case_mode") == "full":
        if not u.get("case_summary"):
            u["case_summary"] = text
        else:
            if not is_same_case(text, u["case_summary"]):
                await message.answer(
                    "Сейчас у вас активна услуга «Полное сопровождение дела» по ранее описанной ситуации.\n\n"
                    "Ваше новое сообщение выглядит как вопрос по другому делу.\n"
                    "В рамках этой услуги я веду только одно дело. Если это связано с тем же делом — поясните, как именно. "
                    "Для нового дела потребуется оформить новую услугу.",
                    reply_markup=kb_main_button(),
                )
                return

        ans = await ask_full_support(text, u.get("case_summary"))
        await message.answer(ans, reply_markup=kb_full_support())
        return

    if u["free_left"] + u["paid_left"] <= 0:
        await message.answer(
            "Ваш бесплатный и оплаченный лимит сообщений сейчас исчерпан.\n\n"
            "Бесплатный лимит обновится через сутки.\n"
            "Если хотите продолжить общение сегодня, вы можете пополнить сообщения в разделе «Пакеты сообщений».",
            reply_markup=kb_main_button(),
        )
        return

    if u["free_left"] > 0:
        u["free_left"] -= 1
    else:
        u["paid_left"] -= 1

    ans = await ask_short_consult(text)
    doc_btn = need_doc_button(ans)

    await message.answer(
        f"{ans}\n\n"
        f"Бесплатных сообщений на сегодня осталось: {u['free_left']}\n"
        f"Оплаченных сообщений: {u['paid_left']}",
        reply_markup=kb_main_button(doc_btn),
    )


async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())


























